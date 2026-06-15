#!/usr/bin/env python3
"""Omni Doc — Razorpay Tech Spec document generator for Nemesis v2.

Creates .docx files matching the Razorpay Tech Spec template format, including:
- Title block with metadata table
- Table of contents
- 16 numbered sections with sub-sections (e.g., 8.1, 8.2, 9.1)
- Styled code blocks, tables, note boxes, bullet/numbered lists
- Approach comparison with Pros/Cons
- Mermaid/sequence diagram placeholders

Reference: https://docs.google.com/document/d/1xbIeTqIGdTIS0CWGdW3dLeydauja4v0Fr0sqKpLe3_I

Usage:
    rubick_doc.py create    --title T --author A --team T [--template tech-spec] --output path
    rubick_doc.py add-section --doc path --number N --content-file F
    rubick_doc.py add-code  --doc path --section N --code-file F [--lang L] [--bg BG]
    rubick_doc.py add-table --doc path --section N --headers "H1,H2,H3" --rows-file F
    rubick_doc.py add-bullet --doc path --text T [--level L]
    rubick_doc.py add-note  --doc path --text T [--label "Note"]
    rubick_doc.py preview   --doc path
    rubick_doc.py finalize  --doc path
"""

import sys
import json
import argparse
import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    import brain_config as cfg
except ImportError:
    cfg = None

# ---------------------------------------------------------------------------
# Template: 16-section Razorpay Tech Spec with sub-sections
# ---------------------------------------------------------------------------

TECH_SPEC_TEMPLATE = [
    {
        "number": 1, "title": "Problem Statement",
        "guidance": "Describe the core problem. What business/technical need is being addressed? Be specific with requirements. Include merchant impact, volume affected, and urgency.",
        "subs": [
            {"number": "1.1", "title": "Business Context", "guidance": "Which merchants, what volume, revenue impact."},
            {"number": "1.2", "title": "Technical Problem", "guidance": "What exactly is broken or missing. Reproduction steps if a bug."},
        ],
    },
    {
        "number": 2, "title": "Introduction & Scope",
        "guidance": "Overview of the system area, what this spec covers, and boundaries.",
        "subs": [
            {"number": "2.1", "title": "Tenets", "guidance": "Security, instrumentation, data handling principles."},
            {"number": "2.2", "title": "Relevant Resources", "guidance": "Links to PRD, design docs, prior art, Slack threads, existing tech specs."},
        ],
    },
    {
        "number": 3, "title": "Out of Scope",
        "guidance": "Explicitly list what is NOT covered. Prevents scope creep. Use a bulleted list with rationale for each exclusion.",
        "subs": [],
    },
    {
        "number": 4, "title": "Futuristic Scope",
        "guidance": "What could this evolve into? Next phases, extensions, related features. Include Phase 2/3 roadmap items.",
        "subs": [],
    },
    {
        "number": 5, "title": "Assumptions, Goals & Non-Goals",
        "guidance": "Split into three clear sub-sections. Each item should be numbered and testable.",
        "subs": [
            {"number": "5.1", "title": "Assumptions", "guidance": "What must be true for this to work. Number each assumption (A1, A2, ...)."},
            {"number": "5.2", "title": "Goals", "guidance": "What success looks like. Measurable outcomes with metrics."},
            {"number": "5.3", "title": "Non-Goals", "guidance": "What we explicitly will NOT do. Prevent scope creep."},
        ],
    },
    {
        "number": 6, "title": "Domain Design",
        "guidance": "Domain model, bounded contexts, key entities, ubiquitous language, business rules, event modeling. Include entity relationship diagram.",
        "subs": [
            {"number": "6.1", "title": "Key Entities & Relationships", "guidance": "Domain entities, their attributes, and relationships. Include ER diagram."},
            {"number": "6.2", "title": "Business Rules & Invariants", "guidance": "Core business rules, amount formulas, validation rules that must always hold."},
            {"number": "6.3", "title": "Ubiquitous Language", "guidance": "Domain terms with precise definitions used throughout the spec."},
        ],
    },
    {
        "number": 7, "title": "Current Architecture / Current HLD",
        "guidance": "How the system works today. Include flow diagrams, sequence diagrams, architecture diagrams. Show what's broken if this is a fix.",
        "subs": [
            {"number": "7.1", "title": "Service Map", "guidance": "All services involved. Canva/Mermaid architecture diagram showing service interactions."},
            {"number": "7.2", "title": "Current Flow", "guidance": "As-is sequence diagram with actual amounts at each step. Show broken behavior if applicable."},
            {"number": "7.3", "title": "Pain Points", "guidance": "What specifically fails in the current architecture. Root cause per service."},
        ],
    },
    {
        "number": 8, "title": "Final Approach - Specifications",
        "guidance": "The chosen solution with FULL detail. This is the CORE section — expect it to be 40-60%% of the doc. Create DEEP sub-sections per flow/feature area (8.1.1, 8.1.1.1 etc). Each sub-section must include: high-level flow diagram, API request/response JSON examples, code diffs, math proofs. For multiple approaches: show each with Pros/Cons table before declaring winner.",
        "subs": [
            {"number": "8.1", "title": "List of Possible Solutions", "guidance": "All approaches considered. For EACH approach: (a) High-Level Description, (b) Architecture Diagram, (c) API Contract Changes with request/response JSON, (d) Detailed Code Changes, (e) Pros (bulleted), (f) Cons (bulleted). End with comparison table and winner declaration."},
            {"number": "8.2", "title": "Chosen Approach — Detailed Specifications", "guidance": "DEEP per-flow breakdown. Create sub-sections per major flow/component (8.2.1 Flow A, 8.2.2 Flow B, etc). Each flow needs: sequence diagram, API examples, code diffs, math proofs. Reference doc example: section 7.1 had 7.1.1 (merchant sourced), 7.1.2 (PG sourced), each with 7.1.x.1 (high level flows)."},
            {"number": "8.3", "title": "Assumptions", "guidance": "Technical assumptions specific to the chosen approach."},
            {"number": "8.4", "title": "Possibility of Open Sourcing", "guidance": "Can this be open sourced? Process."},
            {"number": "8.5", "title": "Possibility on Patent", "guidance": "Does this qualify for a patent?"},
            {"number": "8.6", "title": "Data Model / Schema Changes", "guidance": "FULL schema definitions: CREATE TABLE/ALTER TABLE SQL, struct definitions, proto messages. Include: field types, constraints, indexes, sizing estimates, archival strategy. Use code blocks for schemas."},
            {"number": "8.7", "title": "Business Logic Changes", "guidance": "API endpoint changes with full request/response JSON examples. Pseudocode or real code diffs for logic changes. Flowcharts for complex decision trees. Error states and failure scenarios with HTTP status codes."},
            {"number": "8.8", "title": "Cross-Service Impact", "guidance": "Which services change vs. confirmed safe. Shared contracts (proto, SDK, config). Interface extensions needed. Dependency map diagram."},
            {"number": "8.9", "title": "E2E Flow Trace", "guidance": "Complete end-to-end flow with concrete amounts at every step. Step-by-step table: Step, Service, Code Path, Input, Output, Verification formula."},
            {"number": "8.10", "title": "Miscellaneous Questions", "guidance": "Limitations, failure recovery, future requirements, open items."},
        ],
    },
    {
        "number": 9, "title": "Non-Functional Requirements (NFRs)",
        "guidance": "Address each as applicable. Include specific numbers (TPS, latency p99, uptime %%).",
        "subs": [
            {"number": "9.1", "title": "Scalability", "guidance": "TPS targets, horizontal scaling, bottlenecks. Include load projections."},
            {"number": "9.2", "title": "Availability", "guidance": "Uptime SLA (e.g., 99.99%%), failover, redundancy, circuit breaker config."},
            {"number": "9.3", "title": "Security", "guidance": "No PII in logs, auth mechanisms, credential management, HMAC/signature validation."},
            {"number": "9.4", "title": "Compliance", "guidance": "PCI DSS scope, GDPR data handling, RBI guidelines for payment data."},
            {"number": "9.5", "title": "Reliability", "guidance": "Error budgets, retry strategies with backoff, idempotency keys, timeout cascades."},
            {"number": "9.6", "title": "Infra Cost", "guidance": "Estimated infrastructure cost impact. New pods, memory, CPU, storage."},
        ],
    },
    {
        "number": 10, "title": "Feature Dependencies & SLAs",
        "guidance": "Upstream and downstream applications impacted. Use tables with Service, Impact, SLA, POC columns.",
        "subs": [
            {"number": "10.1", "title": "Upstream Dependencies", "guidance": "Services this feature depends on. Table: Service, Impact, SLA, POC."},
            {"number": "10.2", "title": "Downstream Dependencies", "guidance": "Services impacted by this feature. Note: i18N, Cell Arch, DataHub, Analytics are mandatory downstream checks."},
            {"number": "10.3", "title": "Shared Contracts & Interfaces", "guidance": "Proto files, SDK versions, config keys, feature flags shared across services."},
        ],
    },
    {
        "number": 11, "title": "Testing Plan",
        "guidance": "Comprehensive test strategy. Each test case in a table with: ID, Description, Input, Expected Output, Priority.",
        "subs": [
            {"number": "11.1", "title": "Unit Tests", "guidance": "Per-service unit test cases. Table: Test ID, Service, File, Test Name, What It Verifies."},
            {"number": "11.2", "title": "Integration Tests", "guidance": "Cross-service integration scenarios. Table: Flow, Services, Setup, Expected Outcome."},
            {"number": "11.3", "title": "Regression & Edge Cases", "guidance": "Boundary conditions, backward compatibility, regression matrix. Table: Flow, Touches Changed Code?, Still Works?, Evidence."},
            {"number": "11.4", "title": "UAT Testing", "guidance": "User acceptance test scenarios with real merchant data. Table: #, Scenario, Merchant, Input, Expected Output."},
            {"number": "11.5", "title": "Performance / Load Testing", "guidance": "Load testing plan: TPS targets, soak duration, degradation thresholds."},
        ],
    },
    {
        "number": 12, "title": "Go-Live Plan",
        "guidance": "Production rollout strategy with phased approach.",
        "subs": [
            {"number": "12.1", "title": "Production Rollout & Ramp Plan", "guidance": "Phased rollout: deploy order, feature flags (Splitz/DCS), ramp schedule (1%->5%->25%->50%->100%), monitoring between stages."},
            {"number": "12.2", "title": "Backward Compatibility", "guidance": "Impact on existing consumers. Field additions vs. breaking changes."},
            {"number": "12.3", "title": "Rollback Plan", "guidance": "Per-service rollback procedure. Table: Service, Rollback Method, Time to Effect, Side Effect."},
        ],
    },
    {
        "number": 13, "title": "Monitoring & Logging",
        "guidance": "Metrics, dashboards, alerts, log patterns. Include Grafana panel specs.",
        "subs": [
            {"number": "13.1", "title": "New Metrics", "guidance": "What to instrument. Table: Metric Name, Type (counter/gauge/histogram), Labels, Service."},
            {"number": "13.2", "title": "Dashboards", "guidance": "Grafana dashboard panels needed. Include panel descriptions and queries."},
            {"number": "13.3", "title": "Alert Rules", "guidance": "Thresholds, escalation paths, PagerDuty routing. Table: Alert, Condition, Severity, Runbook."},
            {"number": "13.4", "title": "Log Patterns", "guidance": "Structured log fields to add. grep patterns for troubleshooting."},
        ],
    },
    {
        "number": 14, "title": "Milestones & Timelines",
        "guidance": "Task breakdown with owners. Table: Sr. No, Title, Owner, DevRev/JIRA Link, Due Date, Reviewer, Status.",
        "subs": [
            {"number": "14.1", "title": "Task Breakdown", "guidance": "Detailed task list with effort estimates and dependencies."},
            {"number": "14.2", "title": "Risk Register", "guidance": "Risks with severity, mitigation, owner. Table: Risk, Severity (P0-P3), Mitigation, Owner, Status."},
        ],
    },
    {
        "number": 15, "title": "Glossary",
        "guidance": "Terms and abbreviations used in this document. Table: Term, Definition.",
        "subs": [],
    },
    {
        "number": 16, "title": "Appendix",
        "guidance": "Links to PRDs, milestone trackers, concept notes, project links, @Slash validation logs.",
        "subs": [
            {"number": "16.1", "title": "@Slash Validation Log", "guidance": "Cross-check queries and results from @Slash bot."},
            {"number": "16.2", "title": "References", "guidance": "Links to PRDs, design docs, source code, Slack threads."},
            {"number": "16.3", "title": "Change Log", "guidance": "Document revision history. Table: Version, Date, Author, Changes."},
        ],
    },
]

SOLUTION_DOC_TEMPLATE = [
    {
        "number": 1, "title": "Executive Summary",
        "guidance": "Business impact, scope, timeline. What's broken, which merchants are affected, what this fix achieves.",
        "subs": [],
    },
    {
        "number": 2, "title": "The Problem",
        "guidance": "Bug analysis, reproduction steps, affected payment flows. Include specific failure scenarios.",
        "subs": [
            {"number": "2.1", "title": "Affected Payment Flows", "guidance": "Which flows break and how."},
            {"number": "2.2", "title": "Business Impact", "guidance": "Revenue, merchants, volume affected."},
        ],
    },
    {
        "number": 3, "title": "Domain Model",
        "guidance": "Entity relationships, amount flows, key formulas. Ubiquitous language for the domain.",
        "subs": [
            {"number": "3.1", "title": "Key Entities", "guidance": "DFB, instant discount, fee bearer, etc."},
            {"number": "3.2", "title": "Amount Formula", "guidance": "Core amount relationships."},
        ],
    },
    {
        "number": 4, "title": "Current Architecture",
        "guidance": "Component diagram, data flow, sequence diagrams showing the broken path.",
        "subs": [
            {"number": "4.1", "title": "Service Map", "guidance": "All services involved in the flow."},
            {"number": "4.2", "title": "As-Is Flow", "guidance": "Current broken payment sequence."},
        ],
    },
    {
        "number": 5, "title": "Root Cause Analysis",
        "guidance": "Code-level bug walkthrough with before/after. Trace the exact lines where logic fails.",
        "subs": [
            {"number": "5.1", "title": "Bug Identification", "guidance": "Each bug with file, line, root cause."},
            {"number": "5.2", "title": "Impact Chain", "guidance": "How bugs cascade across services."},
        ],
    },
    {
        "number": 6, "title": "Solution Design",
        "guidance": "Per-fix breakdown with code diffs. Each change: repo, file, what/why/how.",
        "subs": [
            {"number": "6.1", "title": "Change Summary", "guidance": "Table of all changes across repos."},
            {"number": "6.2", "title": "Per-Change Specifications", "guidance": "Code diffs, math proofs, rationale per change."},
            {"number": "6.3", "title": "Alternatives Considered", "guidance": "Approaches rejected and why."},
            {"number": "6.4", "title": "Data Model Changes", "guidance": "Schema/struct/proto modifications."},
        ],
    },
    {
        "number": 7, "title": "Cross-Project Impact",
        "guidance": "Repos changed vs safe, dependency map, shared contracts affected.",
        "subs": [
            {"number": "7.1", "title": "Changed Repositories", "guidance": "Per-repo change list with blast radius."},
            {"number": "7.2", "title": "Unaffected Repositories", "guidance": "Services confirmed safe and why."},
            {"number": "7.3", "title": "Shared Contracts", "guidance": "Proto, SDK, config changes across repos."},
        ],
    },
    {
        "number": 8, "title": "E2E Flow",
        "guidance": "Full payment sequence with fix annotations. Step-by-step amounts at each service hop.",
        "subs": [
            {"number": "8.1", "title": "Running Example", "guidance": "Concrete amounts at each step."},
            {"number": "8.2", "title": "Step-by-Step Trace", "guidance": "Service hop sequence with amounts."},
        ],
    },
    {
        "number": 9, "title": "Feature Flag / DCS Lifecycle",
        "guidance": "Gating strategy, Splitz experiments, DCS config keys, rollout control.",
        "subs": [
            {"number": "9.1", "title": "Splitz Experiments", "guidance": "Feature flag configuration and experiment setup."},
            {"number": "9.2", "title": "DCS Keys", "guidance": "Dynamic config changes required."},
            {"number": "9.3", "title": "Ramp Schedule", "guidance": "Phased rollout percentage plan."},
        ],
    },
    {
        "number": 10, "title": "Deployment Order",
        "guidance": "Ordered deploy steps with DCS gating between phases. Which repo first, which last.",
        "subs": [
            {"number": "10.1", "title": "Deploy Phases", "guidance": "Which repo first, which last, with gates."},
            {"number": "10.2", "title": "Rollback Plan", "guidance": "Per-phase rollback procedure."},
        ],
    },
    {
        "number": 11, "title": "Testing Matrix",
        "guidance": "Per-scenario test plan. Unit tests, integration tests, edge cases per change.",
        "subs": [
            {"number": "11.1", "title": "Unit Tests", "guidance": "Per-change test cases and coverage targets."},
            {"number": "11.2", "title": "Integration Tests", "guidance": "Cross-service test scenarios."},
            {"number": "11.3", "title": "Regression & Edge Cases", "guidance": "Boundary conditions and backward compat."},
            {"number": "11.4", "title": "UAT Plan", "guidance": "User acceptance testing with real merchant data."},
        ],
    },
    {
        "number": 12, "title": "Monitoring & Alerting",
        "guidance": "Metrics, dashboards, alert thresholds, SLOs. What to watch after deploy.",
        "subs": [
            {"number": "12.1", "title": "New Metrics", "guidance": "What to instrument and emit."},
            {"number": "12.2", "title": "Dashboards", "guidance": "Grafana panels and layouts needed."},
            {"number": "12.3", "title": "Alert Rules", "guidance": "Thresholds, escalation paths, runbooks."},
        ],
    },
    {
        "number": 13, "title": "Risk Register",
        "guidance": "Risks with severity, mitigation, status. Blockers and amendments from validation.",
        "subs": [
            {"number": "13.1", "title": "Blockers", "guidance": "Must fix before testing. From risk analysis validation."},
            {"number": "13.2", "title": "Amendments", "guidance": "Required solution changes with effort estimates."},
            {"number": "13.3", "title": "Open Questions", "guidance": "Unresolved items needing investigation."},
        ],
    },
    {
        "number": 14, "title": "Milestones",
        "guidance": "Task breakdown with owners, JIRA/DevRev links, due dates, status.",
        "subs": [],
    },
    {
        "number": 15, "title": "Appendix",
        "guidance": "@Slash validation log, open items, references, links to source docs.",
        "subs": [
            {"number": "15.1", "title": "@Slash Validation Log", "guidance": "Cross-check queries and responses from @Slash bot."},
            {"number": "15.2", "title": "References", "guidance": "Links to source documents, PRDs, design docs."},
            {"number": "15.3", "title": "Change Log", "guidance": "Document revision history."},
        ],
    },
]

TEMPLATES = {
    "tech-spec": TECH_SPEC_TEMPLATE,
    "solution": SOLUTION_DOC_TEMPLATE,
}

# ---------------------------------------------------------------------------
# Style constants matching Razorpay brand
# ---------------------------------------------------------------------------

STYLE = {
    "title_size": 24,
    "title_color": "#0d47a1",
    "section_size": 14,
    "section_color": "#1565c0",
    "sub_section_size": 12,
    "sub_section_color": "#1565c0",
    "sub_header_size": 11.5,
    "sub_header_color": "#37474f",
    "body_size": 10.5,
    "code_size": 9,
    "code_font": "Courier New",
    "code_color": "#37474f",
    "code_bg": "#f5f5f5",
    "code_bg_after": "#e8f5e9",
    "code_bg_before": "#ffebee",
    "table_header_bg": "#1565c0",
    "table_header_color": "#ffffff",
    "meta_key_bg": "#e8eaf6",
    "meta_key_color": "#1a237e",
    "note_bg": "#fff8e1",
    "note_border_color": "#ffc107",
    "note_icon": "\U0001f4a1",
    "pros_bg": "#e8f5e9",
    "cons_bg": "#ffebee",
    "hr_color": "#bdbdbd",
    "guidance_color": "#9e9e9e",
    "margin_top": 2.0,
    "margin_bottom": 2.0,
    "margin_left": 2.5,
    "margin_right": 2.5,
}


# ---------------------------------------------------------------------------
# TechSpecBuilder
# ---------------------------------------------------------------------------

class TechSpecBuilder:
    """Builds a Razorpay Tech Spec .docx document."""

    def __init__(self, title: str, author: str = "Saurav Kumar",
                 team: str = "Payments Core — Emandate / Offers",
                 bu: str = "Payments",
                 template: str = "tech-spec"):
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Run: pip3 install python-docx")

        self.doc = Document()
        self.title = title
        self.author = author
        self.team = team
        self.bu = bu
        self.template_name = template
        self.template = TEMPLATES.get(template, TECH_SPEC_TEMPLATE)
        self._setup_margins()
        self._setup_styles()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin = Cm(STYLE["margin_top"])
            section.bottom_margin = Cm(STYLE["margin_bottom"])
            section.left_margin = Cm(STYLE["margin_left"])
            section.right_margin = Cm(STYLE["margin_right"])

    def _setup_styles(self):
        styles = self.doc.styles
        if "Code Block" not in [s.name for s in styles]:
            style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = STYLE["code_font"]
            style.font.size = Pt(STYLE["code_size"])
            fmt = style.paragraph_format
            fmt.left_indent = Inches(0.4)
            fmt.right_indent = Inches(0.4)
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(4)

    # --- Low-level helpers ---

    @staticmethod
    def _set_color(run, hex_color: str):
        r, g, b = int(hex_color[1:3], 16), int(hex_color[3:5], 16), int(hex_color[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    @staticmethod
    def _shade_cell(cell, fill_hex: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#"))
        tc_pr.append(shd)

    @staticmethod
    def _shade_paragraph(p, fill_hex: str):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#"))
        pPr.append(shd)

    @staticmethod
    def _set_paragraph_border_left(p, color_hex: str, width: int = 12):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(width))
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), color_hex.lstrip("#"))
        pBdr.append(left)
        pPr.append(pBdr)

    # --- Content builders ---

    def add_heading(self, text: str, level: int = 1, color: str = None):
        color = color or STYLE["section_color"]
        p = self.doc.add_heading(text, level=level)
        for run in p.runs:
            self._set_color(run, color)
            run.font.bold = True
        return p

    def add_para(self, text: str = "", bold: bool = False, italic: bool = False,
                 size: float = None, color: str = None, align=None):
        size = size or STYLE["body_size"]
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                self._set_color(run, color)
        return p

    def add_code_block(self, code_text: str, bg: str = None, lang: str = None):
        bg = bg or STYLE["code_bg"]
        if lang:
            label_p = self.doc.add_paragraph()
            label_run = label_p.add_run(lang.upper())
            label_run.bold = True
            label_run.font.size = Pt(8)
            self._set_color(label_run, "#757575")
            label_p.paragraph_format.space_after = Pt(0)

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self._shade_paragraph(p, bg)
        run = p.add_run(code_text)
        run.font.name = STYLE["code_font"]
        run.font.size = Pt(STYLE["code_size"])
        self._set_color(run, STYLE["code_color"])
        return p

    def add_note_box(self, text: str, label: str = "Note"):
        p = self.doc.add_paragraph()
        self._shade_paragraph(p, STYLE["note_bg"])
        self._set_paragraph_border_left(p, STYLE["note_border_color"], width=24)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        icon_run = p.add_run(f"{STYLE['note_icon']} {label}: ")
        icon_run.bold = True
        icon_run.font.size = Pt(STYLE["body_size"])
        text_run = p.add_run(text)
        text_run.font.size = Pt(STYLE["body_size"])
        return p

    def _add_rich_runs(self, p, text: str):
        """Add runs with inline bold (**text**) and code (`text`) formatting."""
        if "**" not in text and "`" not in text:
            run = p.add_run(text)
            run.font.size = Pt(STYLE["body_size"])
            return
        tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
                run.font.size = Pt(STYLE["body_size"])
            elif token.startswith("`") and token.endswith("`"):
                run = p.add_run(token[1:-1])
                run.font.name = STYLE["code_font"]
                run.font.size = Pt(STYLE["code_size"])
            else:
                run = p.add_run(token)
                run.font.size = Pt(STYLE["body_size"])

    def add_bullet(self, text: str, level: int = 0):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
        self._add_rich_runs(p, text)
        return p

    def add_numbered(self, text: str, level: int = 0):
        p = self.doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
        self._add_rich_runs(p, text)
        return p

    def add_table(self, headers: list, rows: list, col_widths: list = None):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr_row = t.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    self._set_color(run, STYLE["table_header_color"])
            self._shade_cell(cell, STYLE["table_header_bg"])

        for ri, row_data in enumerate(rows):
            row = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                if ci < len(row.cells):
                    cell = row.cells[ci]
                    val_str = str(val)
                    if "**" in val_str or "`" in val_str:
                        cell.text = ""
                        p = cell.paragraphs[0]
                        tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", val_str)
                        for token in tokens:
                            if token.startswith("**") and token.endswith("**"):
                                run = p.add_run(token[2:-2])
                                run.bold = True
                                run.font.size = Pt(10)
                            elif token.startswith("`") and token.endswith("`"):
                                run = p.add_run(token[1:-1])
                                run.font.name = STYLE["code_font"]
                                run.font.size = Pt(9)
                            else:
                                run = p.add_run(token)
                                run.font.size = Pt(10)
                    else:
                        cell.text = val_str
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(10)

        if col_widths:
            for ci, w in enumerate(col_widths):
                for row in t.rows:
                    if ci < len(row.cells):
                        row.cells[ci].width = Inches(w)
        return t

    def add_pros_cons(self, pros: list, cons: list):
        self.add_para("Pros", bold=True, size=11, color="#2e7d32")
        for item in pros:
            p = self.add_bullet(item)
            self._shade_paragraph(p, STYLE["pros_bg"])

        self.add_para("Cons", bold=True, size=11, color="#c62828")
        for item in cons:
            p = self.add_bullet(item)
            self._shade_paragraph(p, STYLE["cons_bg"])

    def add_diagram_placeholder(self, title: str, diagram_type: str = "sequence"):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_paragraph(p, "#e3f2fd")
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"[{diagram_type.upper()} DIAGRAM: {title}]")
        run.italic = True
        run.font.size = Pt(10)
        self._set_color(run, "#1565c0")
        return p

    def add_image(self, path: str, width: float = 5.5, caption: str = None):
        """Embed a local image file (PNG/JPG) into the document."""
        if not os.path.exists(path):
            return self.add_diagram_placeholder(caption or path, "image")
        self.doc.add_picture(path, width=Inches(width))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            self._set_color(run, "#757575")
        return last_paragraph

    def add_hr(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), STYLE["hr_color"].lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # --- Section builders ---

    def section_header(self, num, title: str):
        self.add_hr()
        h = self.add_heading(f"{num}. {title}", level=1, color=STYLE["section_color"])
        h.runs[0].font.size = Pt(STYLE["section_size"])
        return h

    def sub_section_header(self, num: str, title: str):
        text = f"{num}. {title}" if num else title
        h = self.add_heading(text, level=2, color=STYLE["sub_section_color"])
        h.runs[0].font.size = Pt(STYLE["sub_section_size"])
        return h

    def sub_header(self, title: str, color: str = None):
        color = color or STYLE["sub_header_color"]
        h = self.add_heading(title, level=3, color=color)
        h.runs[0].font.size = Pt(STYLE["sub_header_size"])
        return h

    # --- Document builders ---

    def build_title_block(self):
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"Tech Spec - {self.title}")
        title_run.bold = True
        title_run.font.size = Pt(STYLE["title_size"])
        self._set_color(title_run, STYLE["title_color"])

        self.doc.add_paragraph()

        today = datetime.now().strftime("%Y-%m-%d")
        meta_data = [
            ("Author", self.author),
            ("Team / Pod", f"{self.team} | BU: {self.bu}"),
            ("Published Date", today),
            ("Reviewer Name", ""),
            ("Reviewed Date", ""),
            ("Status", "Draft"),
        ]

        meta = self.doc.add_table(rows=len(meta_data), cols=2)
        meta.style = "Table Grid"
        for i, (k, v) in enumerate(meta_data):
            row = meta.rows[i]
            key_cell = row.cells[0]
            val_cell = row.cells[1]
            key_cell.text = k
            val_cell.text = v
            for p in key_cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    self._set_color(run, STYLE["meta_key_color"])
            for p in val_cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            self._shade_cell(key_cell, STYLE["meta_key_bg"])

        meta.columns[0].width = Inches(1.8)
        meta.columns[1].width = Inches(4.2)
        self.doc.add_paragraph()

        # Team label
        team_p = self.doc.add_paragraph()
        team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        team_run = team_p.add_run("Payments Core Team")
        team_run.bold = True
        team_run.font.size = Pt(11)
        self._set_color(team_run, STYLE["section_color"])

        self.doc.add_paragraph()

    def build_toc(self):
        toc_h = self.add_heading("Table of Contents", level=1, color=STYLE["section_color"])
        toc_h.runs[0].font.size = Pt(STYLE["section_size"])

        for sec in self.template:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{sec['number']}. {sec['title']}")
            run.bold = True
            run.font.size = Pt(STYLE["body_size"])
            self._set_color(run, STYLE["section_color"])
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)

            for sub in sec.get("subs", []):
                sp = self.doc.add_paragraph()
                sp.paragraph_format.left_indent = Inches(0.4)
                sp.paragraph_format.space_after = Pt(1)
                sp.paragraph_format.space_before = Pt(0)
                sr = sp.add_run(f"{sub['number']}. {sub['title']}")
                sr.font.size = Pt(9.5)
                self._set_color(sr, "#546e7a")

        self.doc.add_paragraph()

    def build_note_box(self):
        self.add_note_box(
            "Read following docs to get expert on different topics - Full Reading List - Link",
            label="Note"
        )
        self.doc.add_paragraph()

    def build_skeleton(self):
        self.build_title_block()
        self.build_toc()
        self.build_note_box()

        for sec in self.template:
            self.section_header(sec["number"], sec["title"])
            self.add_para(
                f'[{sec["guidance"]}]',
                italic=True, color=STYLE["guidance_color"], size=10
            )

            for sub in sec.get("subs", []):
                self.sub_section_header(sub["number"], sub["title"])
                self.add_para(
                    f'[{sub["guidance"]}]',
                    italic=True, color=STYLE["guidance_color"], size=10
                )

            # Add pre-built tables for specific sections
            if sec["number"] == 10:
                for sub in sec.get("subs", []):
                    pass  # guidance already added above
                self._add_dependency_table_stub()
            elif sec["number"] == 14:
                self._add_milestone_table_stub()

    def _add_dependency_table_stub(self):
        self.add_table(
            headers=["Sr. No.", "Service", "Impact", "SLA (if applicable)", "POC for Service"],
            rows=[["1", "", "", "", ""]],
        )
        self.doc.add_paragraph()

    def _add_milestone_table_stub(self):
        self.add_table(
            headers=["Sr. No.", "Title", "Owner", "JIRA/DevRev Link", "Due Date", "Reviewer", "Status"],
            rows=[["1", "Milestone 1", "", "", "", "", ""]],
        )
        self.doc.add_paragraph()

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path

    def get_structure(self) -> list[dict]:
        result = []
        current_section = None
        para_count = 0
        table_count = 0
        code_count = 0

        def _is_heading(element):
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    val = pStyle.get(qn("w:val")) or ""
                    return val.startswith("Heading")
            return False

        for element in self.doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                text = ""
                for child in element.iter():
                    t = child.text or ""
                    if t.strip():
                        text = t.strip()
                        break

                is_heading = _is_heading(element)

                if is_heading and text and re.match(r"^\d+\.\s", text):
                    if current_section:
                        result.append({
                            "number": current_section["number"],
                            "title": current_section["title"],
                            "paragraphs": para_count,
                            "tables": table_count,
                            "code_blocks": code_count,
                            "status": "filled" if para_count > 1 or table_count > 0 else "skeleton",
                        })
                    match = re.match(r"^(\d+)\.\s(.+)", text)
                    if match:
                        current_section = {
                            "number": int(match.group(1)),
                            "title": match.group(2),
                        }
                        para_count = 0
                        table_count = 0
                        code_count = 0
                elif is_heading and text and re.match(r"^\d+\.\d+\.\s", text):
                    para_count += 1
                else:
                    is_code = False
                    for child in element.iter():
                        rPr = child.find(qn("w:rPr"))
                        if rPr is not None:
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is not None and "Courier" in (rFonts.get(qn("w:ascii")) or ""):
                                is_code = True
                                break
                    if is_code:
                        code_count += 1
                    else:
                        para_count += 1
            elif tag == "tbl":
                table_count += 1

        if current_section:
            result.append({
                "number": current_section["number"],
                "title": current_section["title"],
                "paragraphs": para_count,
                "tables": table_count,
                "code_blocks": code_count,
                "status": "filled" if para_count > 1 or table_count > 0 else "skeleton",
            })

        return result

    @staticmethod
    def detect_template(doc: Document) -> list[dict]:
        """Detect which template a document was built with by inspecting section 1's title."""
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag != "p":
                continue
            text = ""
            for child in element.iter():
                t = child.text or ""
                if t.strip():
                    text = t.strip()
                    break
            match = re.match(r"^1\.\s(.+)", text)
            if match:
                title = match.group(1).strip()
                if title == "Executive Summary":
                    return SOLUTION_DOC_TEMPLATE
                return TECH_SPEC_TEMPLATE
        return TECH_SPEC_TEMPLATE


# ---------------------------------------------------------------------------
# Markdown-to-docx converter
# ---------------------------------------------------------------------------

def _md_to_docx(builder: TechSpecBuilder, content: str):
    lines = content.strip().split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = None

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith("```"):
            if in_code_block:
                builder.add_code_block("\n".join(code_lines), lang=code_lang)
                code_lines = []
                code_lang = None
                in_code_block = False
            else:
                in_code_block = True
                lang_match = re.match(r"```(\w+)", line.strip())
                code_lang = lang_match.group(1) if lang_match else None
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            builder.sub_header(stripped[5:])
        elif stripped.startswith("### "):
            builder.sub_header(stripped[4:])
        elif stripped.startswith("## "):
            sub_text = stripped[3:]
            sub_match = re.match(r"^(\d+\.\d+)\.\s+(.+)", sub_text)
            if sub_match:
                builder.sub_section_header(sub_match.group(1), sub_match.group(2))
            else:
                builder.sub_section_header("", sub_text)
        # Note box
        elif stripped.startswith("> **Note"):
            text = stripped.lstrip("> ").replace("**Note**: ", "").replace("**Note:** ", "")
            builder.add_note_box(text)
        elif stripped.startswith("> "):
            text = stripped[2:]
            builder.add_note_box(text)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            builder.add_bullet(stripped[2:])
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            builder.add_numbered(re.sub(r"^\d+\.\s", "", stripped))
        # Table row (skip separator rows)
        elif stripped.startswith("|") and not re.match(r"^\|[\s\-|:]+\|$", stripped):
            # Collect full table
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1

            # Parse table
            data_rows = [l for l in table_lines if not re.match(r"^\|[\s\-|:]+\|$", l)]
            if data_rows:
                headers = [c.strip() for c in data_rows[0].split("|")[1:-1]]
                rows = []
                for dr in data_rows[1:]:
                    cells = [c.strip() for c in dr.split("|")[1:-1]]
                    rows.append(cells)
                if headers:
                    builder.add_table(headers, rows)

            i = j
            continue
        # Pros/Cons block
        elif stripped.replace("*", "").replace(":", "").strip().lower() == "pros":
            pros = []
            j = i + 1
            while j < len(lines) and (lines[j].strip().startswith("- ") or lines[j].strip().startswith("* ")):
                pros.append(lines[j].strip()[2:])
                j += 1
            cons = []
            if j < len(lines) and lines[j].strip().replace("*", "").replace(":", "").strip().lower() == "cons":
                j += 1
                while j < len(lines) and (lines[j].strip().startswith("- ") or lines[j].strip().startswith("* ")):
                    cons.append(lines[j].strip()[2:])
                    j += 1
            if pros or cons:
                builder.add_pros_cons(pros, cons)
            i = j
            continue
        # Image embedding: ![caption](path)
        elif re.match(r'^!\[.*\]\(.*\)$', stripped):
            img_match = re.match(r'^!\[(.*)\]\((.*)\)$', stripped)
            if img_match:
                caption, img_path = img_match.group(1), img_match.group(2)
                if os.path.exists(img_path):
                    builder.add_image(img_path, width=5.5, caption=caption or None)
                else:
                    builder.add_diagram_placeholder(caption or img_path, "image")
        # Diagram placeholder
        elif re.match(r"^\[.*(diagram|sequence|flow).*\]$", stripped, re.IGNORECASE):
            inner = stripped[1:-1]
            parts = inner.split(":", 1)
            if len(parts) == 2:
                builder.add_diagram_placeholder(parts[1].strip(), parts[0].strip().lower())
            else:
                builder.add_diagram_placeholder(inner, "diagram")
        # Regular paragraph
        else:
            if "**" in stripped or "`" in stripped:
                p = builder.doc.add_paragraph()
                builder._add_rich_runs(p, stripped)
            else:
                builder.add_para(stripped)

        i += 1


# ---------------------------------------------------------------------------
# CLI commands
# ---------------------------------------------------------------------------

def cmd_create(args):
    builder = TechSpecBuilder(
        title=args.title,
        author=args.author,
        team=args.team,
        bu=args.bu,
        template=args.template,
    )
    builder.build_skeleton()
    path = builder.save(args.output)
    size = os.path.getsize(path)
    sections = len(builder.template)
    sub_sections = sum(len(s.get("subs", [])) for s in builder.template)
    print(json.dumps({
        "ok": True,
        "path": path,
        "size_bytes": size,
        "template": args.template,
        "sections": sections,
        "sub_sections": sub_sections,
    }))


def _find_section_range(body, section_num):
    """Find the element range for a given section number.

    Returns (heading_idx, end_idx) where:
    - heading_idx: index of the section Heading 1 element
    - end_idx: index of the HR element preceding the next section (or len(body))
    """
    elements = list(body)
    heading_idx = None
    end_idx = len(elements)

    for i, el in enumerate(elements):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag != "p":
            continue
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None or not (pStyle.get(qn("w:val")) or "").startswith("Heading1"):
            continue
        text = ""
        for child in el.iter():
            t = child.text or ""
            if t.strip():
                text = t.strip()
                break
        match = re.match(r"^(\d+)\.\s", text)
        if not match:
            continue
        num = int(match.group(1))
        if num == section_num:
            heading_idx = i
        elif heading_idx is not None and num > section_num:
            # The HR paragraph before this heading belongs to the next section
            # Check if elements[i-1] is an HR paragraph
            if i > 0:
                prev = elements[i - 1]
                ptag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
                if ptag == "p":
                    ppPr = prev.find(qn("w:pPr"))
                    if ppPr is not None and ppPr.find(qn("w:pBdr")) is not None:
                        end_idx = i - 1
                    else:
                        end_idx = i
                else:
                    end_idx = i
            else:
                end_idx = i
            break

    return heading_idx, end_idx


def cmd_add_section(args):
    doc = Document(args.doc)
    content = ""
    if args.content_file:
        with open(args.content_file, "r") as f:
            content = f.read()
    elif args.content:
        content = args.content
    else:
        print(json.dumps({"error": "No content provided. Use --content-file or --content"}))
        sys.exit(1)

    body = doc.element.body
    heading_idx, end_idx = _find_section_range(body, args.number)
    if heading_idx is None:
        print(json.dumps({"error": f"Section {args.number} not found in document"}))
        sys.exit(1)

    elements = list(body)
    for el in elements[heading_idx + 1:end_idx]:
        body.remove(el)

    temp_doc = Document()
    temp_builder = TechSpecBuilder.__new__(TechSpecBuilder)
    temp_builder.doc = temp_doc
    temp_builder.template = TECH_SPEC_TEMPLATE
    _md_to_docx(temp_builder, content)

    insert_after = elements[heading_idx]
    for el in list(temp_doc.element.body):
        insert_after.addnext(el)
        insert_after = el

    doc.save(args.doc)
    print(json.dumps({"ok": True, "section": args.number, "doc": args.doc}))


def cmd_add_code(args):
    doc = Document(args.doc)
    code = ""
    if args.code_file:
        with open(args.code_file, "r") as f:
            code = f.read()
    elif args.code:
        code = args.code
    else:
        print(json.dumps({"error": "No code provided"}))
        sys.exit(1)

    builder = TechSpecBuilder.__new__(TechSpecBuilder)
    builder.doc = doc
    bg = args.bg or STYLE["code_bg"]
    builder.add_code_block(code, bg=bg, lang=args.lang)
    doc.save(args.doc)
    print(json.dumps({"ok": True, "section": args.section, "doc": args.doc}))


def cmd_add_table(args):
    doc = Document(args.doc)
    headers = [h.strip() for h in args.headers.split(",")]

    rows = []
    if args.rows_file:
        with open(args.rows_file, "r") as f:
            data = json.load(f)
            rows = data.get("rows", data) if isinstance(data, dict) else data
    elif args.rows:
        data = json.loads(args.rows)
        rows = data.get("rows", data) if isinstance(data, dict) else data

    builder = TechSpecBuilder.__new__(TechSpecBuilder)
    builder.doc = doc
    builder.add_table(headers, rows)
    doc.save(args.doc)
    print(json.dumps({"ok": True, "section": args.section, "doc": args.doc,
                       "rows": len(rows), "cols": len(headers)}))


def cmd_add_bullet(args):
    doc = Document(args.doc)
    builder = TechSpecBuilder.__new__(TechSpecBuilder)
    builder.doc = doc
    builder.add_bullet(args.text, level=args.level)
    doc.save(args.doc)
    print(json.dumps({"ok": True, "doc": args.doc}))


def cmd_add_note(args):
    doc = Document(args.doc)
    builder = TechSpecBuilder.__new__(TechSpecBuilder)
    builder.doc = doc
    builder.add_note_box(args.text, label=args.label)
    doc.save(args.doc)
    print(json.dumps({"ok": True, "doc": args.doc}))


def cmd_preview(args):
    doc = Document(args.doc)
    builder = TechSpecBuilder.__new__(TechSpecBuilder)
    builder.doc = doc
    if args.template:
        builder.template = TEMPLATES[args.template]
    else:
        builder.template = TechSpecBuilder.detect_template(doc)
    structure = builder.get_structure()

    template_name = args.template or (
        "solution" if builder.template is SOLUTION_DOC_TEMPLATE else "tech-spec"
    )
    filled = sum(1 for s in structure if s["status"] == "filled")
    total = len(structure)
    print(json.dumps({
        "doc": args.doc,
        "template": template_name,
        "sections": structure,
        "total": total,
        "filled": filled,
        "completion": f"{filled}/{total}",
    }, indent=2))


def cmd_embed_images(args):
    """Embed images into a doc post-section-fill. Avoids cross-doc rId bug.

    The add-section command uses a temp Document to convert markdown, then copies
    XML elements to the real doc. Images added to the temp doc create rId references
    that break when elements are moved. This command adds images DIRECTLY to the
    real doc, bypassing the cross-document relationship issue.
    """
    doc = Document(args.doc)
    with open(args.image_map, "r") as f:
        image_map = json.load(f)
        # Expected format: list of {"section": N, "path": "...", "caption": "...", "position": "start"|"end"}

    if isinstance(image_map, dict):
        image_map = list(image_map.values())

    embedded = 0
    for info in image_map:
        img_path = info.get("path", "")
        caption = info.get("caption", "")
        section_num = info.get("section")
        position = info.get("position", "end")

        if not img_path or not os.path.exists(img_path):
            continue
        if section_num is None:
            continue

        heading_idx, end_idx = _find_section_range(doc.element.body, int(section_num))
        if heading_idx is None:
            continue

        elements = list(doc.element.body)

        # Determine insertion point
        if position == "start":
            anchor = elements[heading_idx]
        else:
            anchor = elements[min(end_idx - 1, len(elements) - 1)]

        # Add picture directly to the real doc (creates proper rId relationship)
        doc.add_picture(img_path, width=Inches(5.5))
        pic_p = doc.paragraphs[-1]
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Move the picture paragraph to the correct position
        anchor.addnext(pic_p._p)

        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            TechSpecBuilder._set_color(run, "#757575")
            # Insert caption right after the picture
            pic_p._p.addnext(cap_p._p)

        embedded += 1

    doc.save(args.doc)
    size = os.path.getsize(args.doc)
    print(json.dumps({"ok": True, "images_embedded": embedded,
                       "doc": args.doc, "size_bytes": size}))


def cmd_finalize(args):
    doc = Document(args.doc)

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    doc.save(args.doc)
    size = os.path.getsize(args.doc)
    print(json.dumps({"ok": True, "path": args.doc, "size_bytes": size,
                       "status": "finalized", "page_numbers": True}))


def main():
    if not HAS_DOCX:
        print(json.dumps({"error": "python-docx not installed. Run: pip3 install python-docx"}))
        sys.exit(1)

    parser = argparse.ArgumentParser(description="Omni Doc — Tech Spec Generator")
    sub = parser.add_subparsers(dest="command")

    p_create = sub.add_parser("create")
    p_create.add_argument("--title", required=True)
    p_create.add_argument("--author", default="Saurav Kumar")
    p_create.add_argument("--team", default="Payments Core — Emandate / Offers")
    p_create.add_argument("--bu", default="Payments")
    p_create.add_argument("--template", default="tech-spec", choices=list(TEMPLATES.keys()))
    p_create.add_argument("--output", required=True)

    p_section = sub.add_parser("add-section")
    p_section.add_argument("--doc", required=True)
    p_section.add_argument("--number", type=int, required=True)
    p_section.add_argument("--content-file", default=None)
    p_section.add_argument("--content", default=None)

    p_code = sub.add_parser("add-code")
    p_code.add_argument("--doc", required=True)
    p_code.add_argument("--section", type=int, required=True)
    p_code.add_argument("--code-file", default=None)
    p_code.add_argument("--code", default=None)
    p_code.add_argument("--lang", default=None)
    p_code.add_argument("--bg", default=None)

    p_table = sub.add_parser("add-table")
    p_table.add_argument("--doc", required=True)
    p_table.add_argument("--section", type=int, required=True)
    p_table.add_argument("--headers", required=True)
    p_table.add_argument("--rows-file", default=None)
    p_table.add_argument("--rows", default=None)

    p_bullet = sub.add_parser("add-bullet")
    p_bullet.add_argument("--doc", required=True)
    p_bullet.add_argument("--text", required=True)
    p_bullet.add_argument("--level", type=int, default=0)

    p_note = sub.add_parser("add-note")
    p_note.add_argument("--doc", required=True)
    p_note.add_argument("--text", required=True)
    p_note.add_argument("--label", default="Note")

    p_preview = sub.add_parser("preview")
    p_preview.add_argument("--doc", required=True)
    p_preview.add_argument("--template", default=None, choices=list(TEMPLATES.keys()))

    p_embed = sub.add_parser("embed-images")
    p_embed.add_argument("--doc", required=True)
    p_embed.add_argument("--image-map", required=True, help="JSON file with image specs")

    p_final = sub.add_parser("finalize")
    p_final.add_argument("--doc", required=True)

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    cmds = {
        "create": cmd_create,
        "add-section": cmd_add_section,
        "add-code": cmd_add_code,
        "add-table": cmd_add_table,
        "add-bullet": cmd_add_bullet,
        "add-note": cmd_add_note,
        "embed-images": cmd_embed_images,
        "preview": cmd_preview,
        "finalize": cmd_finalize,
    }

    cmds[args.command](args)


if __name__ == "__main__":
    main()
